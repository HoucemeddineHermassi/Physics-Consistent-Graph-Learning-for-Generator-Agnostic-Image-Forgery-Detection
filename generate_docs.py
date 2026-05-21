import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_cover_letter():
    doc = Document()
    
    # Set font to standard professional (Times New Roman or Calibri)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Correspondence Information
    header = doc.add_paragraph()
    header.add_run("Houcemeddine HERMASSI").bold = True
    header.add_run("\nLR-RISC ENIT, University of Elmanar / ENICarthage, University of Carthage")
    header.add_run("\nTunis, Tunisia")
    header.add_run("\nEmail: houcemeddine.hermassi@enit.rnu")
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))

    # Recipient
    doc.add_paragraph("\nEditors-in-Chief\nInformation Sciences\nElsevier")

    # Subject
    doc.add_paragraph("\nDear Editors-in-Chief,\n")

    # Body
    body = doc.add_paragraph()
    body.add_run("We are pleased to submit our original research manuscript entitled ")
    body.add_run("\"Causal Physical Consistency Learning for Blind Image Forgery Detection\"").italic = True
    body.add_run(" for consideration for publication as a Research Article in ")
    body.add_run("Information Sciences.").bold = True

    p = doc.add_paragraph()
    p.add_run("In the current era of photorealistic AI-driven image synthesis, traditional forensic methods relying on low-level statistical artifacts are increasingly failing. This paper introduces a paradigm shift by moving from artifact recognition to physical causality verification. Our proposed framework, Causal Physical Consistency Learning (CPCL), detects forgeries by identifying 'causal breaks' in the underlying physical laws of image formation (illumination, reflectance, and geometry).")

    p = doc.add_paragraph("Key innovations presented in this work include:")
    list_items = [
        "A self-supervised Physics Proxy based on TV-Retinex decomposition that extracts intrinsic features invariant to generative architectures.",
        "A Causal Graph Attention Network (CGAT) that evaluates local physical consistency across image patches.",
        "The introduction of the PhysForge benchmark, a controlled dataset containing 15,000 forgeries specifically designed to isolate physics-level violations.",
        "Significant performance improvements (+19% cross-dataset F1-score) over state-of-the-art detectors on diffusion-based forgeries (e.g., CocoGlide)."
    ]
    for item in list_items:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run("Given the journal's focus on ")
    p.add_run("Information Sciences").italic = True
    p.add_run("' interest in robust AI, data security, and forensic analysis, we believe our work offers a timely and significant contribution to the field of digital trustworthy systems. This manuscript has not been published elsewhere and is not under consideration by another journal. All authors have approved the final manuscript.")

    doc.add_paragraph("\nThank you for your time and consideration.\n\nSincerely,\n\nHoucemeddine HERMASSI")

    # Save
    doc.save("Cover_Letter_Information_Sciences.docx")
    print("Cover letter generated: Cover_Letter_Information_Sciences.docx")

def create_readme():
    content = """# CPCL-Forgeries: Causal Physical Consistency Learning

[![Paper](https://img.shields.io/badge/Journal-Information%20Sciences-blue)](https://www.sciencedirect.com/journal/information-sciences)
[![PyTorch](https://img.shields.io/badge/Framework-PyTorch-ee4c2c)](https://pytorch.org/)

Official implementation of the paper: **"Causal Physical Consistency Learning for Blind Image Forgery Detection"**.

## Overview
As generative image synthesis reaches photorealistic fidelity, traditional forensics methods relying on textural artifacts are increasingly susceptible to failures. **CPCL** is a paradigm-shifting approach that detects forgeries by verifying the underlying **physical causality** of image formation. 

Instead of searching for generator-specific fingerprints, we model authenticity as the adherence to a globally consistent physical world model where illumination, reflectance, and geometry are causally linked.

### Key Contributions
- **Causal Forensic Paradigm**: Framing forgery as a local causal intervention ($do$-calculus) in the physical image formation process.
- **Physics-First Architecture**: Lightweight framework integrating a self-supervised physics proxy with a Causal Graph Attention Network (CGAT).
- **PhysForge Benchmark**: A controlled dataset of 15,000 images specifically designed to isolate lighting, shadow, and reflectance violations.

---

## Getting Started

### Prerequisites
- Python 3.8+
- PyTorch 1.12+
- LaTeX environment (e.g., MiKTeX) to compile papers.

### Installation
```bash
git clone https://github.com/YourUsername/CPCL-Forgeries.git
cd CPCL-Forgeries
pip install -r requirements.txt
```

---

## Repository Structure
- `models/`: Implementation of the Physics Proxy and CGAT architecture.
- `data_loader.py`: Integration for CASIA v2.0, CocoGlide, Columbia, and In-the-Wild datasets.
- `run_simulation.py`: Main script for the comprehensive experimental pipeline.
- `els-cas-templates/`: LaTeX source and figures for the manuscript.
- `generate_*.py`: Set of scripts to reproduce all figures and tables in the paper.

---

## Reproducibility
We provide dedicated scripts to reproduce every figure presented in the manuscript:

| Figure | Script | Description |
| :--- | :--- | :--- |
| **Fig. 1** | `generate_scm_diagram.py` | SCM Causal Graph visualization. |
| **Fig. 4** | `generate_attention_maps.py` | Qualitative qualitative maps. |
| **Fig. 6** | `generate_attention_maps.py` | Physical "Causal Break" visualization. |
| **Fig. 7-9** | `generate_robustness_plots.py` | Robustness curves (JPEG, Blur, etc.). |
| **Fig. 11** | `generate_calibration_plot.py` | Reliability diagrams (MC Dropout). |
| **Fig. 13** | `generate_explainability_viz.py` | Mechanistic interpretability analysis. |
| **Fig. 14** | `generate_scalability_plot.py` | Resource efficiency analysis. |

To run the full simulation:
```bash
python run_simulation.py
```

---

## Dataset: PhysForge
The **PhysForge** benchmark is included in this repository. It provides a targeted evaluation for:
1. **Lighting Mismatch**
2. **Shadow Inconsistency**
3. **Reflectance Anomaly**

---

## Citation
If you find this work useful, please cite our paper:
```bibtex
@article{hermassi2024cpcl,
  title={Causal Physical Consistency Learning for Blind Image Forgery Detection},
  author={Hermassi, Houcemeddine},
  journal={Information Sciences},
  volume={--},
  year={2024},
  publisher={Elsevier}
}
```
"""
    with open("README.md", "w", encoding="utf-8") as f:
        f.write(content)
    print("README.md generated.")

if __name__ == "__main__":
    create_cover_letter()
    create_readme()
